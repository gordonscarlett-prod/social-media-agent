"""
Google Drive integration — browse local Social Media folder and get public shareable URLs.
Files live locally at SOCIAL_MEDIA_DIR; when selected for a post they are uploaded to
Google Drive and a public direct-download URL is returned.
"""
import os
import json
import mimetypes
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

SOCIAL_MEDIA_DIR = os.getenv(
    "SOCIAL_MEDIA_DIR",
    r"C:\Users\HP\OneDrive\MenardProperties\Social Media",
)
# The Drive folder the user publishes media from: https://drive.google.com/drive/folders/<id>
SOCIAL_MEDIA_FOLDER_ID = os.getenv(
    "SOCIAL_MEDIA_FOLDER_ID",
    "1dI95ZeFBW_WjFscAfh6aR__BG7G5Eiyu",
)
CREDENTIALS_FILE = os.getenv("GOOGLE_CREDENTIALS_FILE", "google_credentials.json")
TOKEN_FILE = os.getenv("GOOGLE_TOKEN_FILE", "google_token.json")
CODE_VERIFIER_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "google_code_verifier.json")

# drive.file      -> create/manage files the app uploads (used for Media Library "Get URL")
# drive.readonly  -> browse/read any folder shared with the user (used for Resource Library)
SCOPES = [
    "https://www.googleapis.com/auth/drive.file",
    "https://www.googleapis.com/auth/drive.readonly",
]

GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
GOOGLE_FOLDER_MIME = "application/vnd.google-apps.folder"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME = "application/pdf"

IMAGE_EXTS  = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".heic", ".svg"}
VIDEO_EXTS  = {".mp4", ".mov", ".avi", ".webm", ".mkv"}
FOLDER_TYPE_MAP = {
    "Images":       "image",
    "Infographics": "image",
    "Reels":        "video",
    "Videos":       "video",
    "Blog":         "document",
}


# ─────────────────────────────────────────────
# Local folder browsing
# ─────────────────────────────────────────────

def list_folder(relative_path: str = "") -> dict:
    """
    List files and subfolders under SOCIAL_MEDIA_DIR / relative_path.
    Returns { folders: [...], files: [...] }
    """
    base = Path(SOCIAL_MEDIA_DIR)
    target = (base / relative_path).resolve()

    # Safety: never escape the Social Media root
    if not str(target).startswith(str(base)):
        raise PermissionError("Path outside Social Media directory")

    if not target.exists():
        raise FileNotFoundError(f"Folder not found: {target}")

    folders = []
    files   = []

    for item in sorted(target.iterdir()):
        if item.name.startswith("."):
            continue
        if item.is_dir():
            folders.append({
                "name": item.name,
                "path": str(item.relative_to(base)).replace("\\", "/"),
            })
        elif item.is_file():
            ext = item.suffix.lower()
            media_type = (
                "image"    if ext in IMAGE_EXTS  else
                "video"    if ext in VIDEO_EXTS  else
                "document"
            )
            files.append({
                "name":       item.name,
                "path":       str(item.relative_to(base)).replace("\\", "/"),
                "full_path":  str(item),
                "size":       item.stat().st_size,
                "media_type": media_type,
                "ext":        ext,
            })

    return {"folders": folders, "files": files, "current_path": relative_path}


# ─────────────────────────────────────────────
# Google Drive OAuth helpers
# ─────────────────────────────────────────────

def _token_has_required_scopes() -> bool:
    """Check the raw token file (not the SDK default) for the scopes we currently need."""
    if not os.path.exists(TOKEN_FILE):
        return False
    try:
        with open(TOKEN_FILE) as f:
            info = json.load(f)
        granted = set(info.get("scopes", []))
        return set(SCOPES).issubset(granted)
    except Exception:
        return False


def _get_credentials():
    """Load or refresh OAuth2 credentials (token only — no interactive flow here)."""
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request

    creds = None

    if os.path.exists(TOKEN_FILE) and _token_has_required_scopes():
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            try:
                creds.refresh(Request())
                with open(TOKEN_FILE, "w") as f:
                    f.write(creds.to_json())
            except Exception:
                creds = None
                if os.path.exists(TOKEN_FILE):
                    os.remove(TOKEN_FILE)

        if not creds or not creds.valid:
            raise RuntimeError(
                "Google Drive is not authenticated. "
                "Visit /api/drive/auth in your browser to authorize."
            )

    return creds


def get_auth_url(redirect_uri: str) -> str:
    """Generate the Google OAuth2 authorization URL for the web-based flow."""
    from google_auth_oauthlib.flow import Flow

    if not os.path.exists(CREDENTIALS_FILE):
        raise FileNotFoundError(
            f"Google credentials file not found: {CREDENTIALS_FILE}. "
            "Download it from Google Cloud Console → APIs & Services → Credentials."
        )

    flow = Flow.from_client_secrets_file(
        CREDENTIALS_FILE,
        scopes=SCOPES,
        redirect_uri=redirect_uri,
    )
    auth_url, _ = flow.authorization_url(
        access_type="offline",
        include_granted_scopes="true",
        prompt="consent",
    )
    # Persist PKCE code_verifier so exchange_auth_code can use it
    code_verifier = getattr(flow, "code_verifier", None)
    if code_verifier:
        with open(CODE_VERIFIER_FILE, "w") as f:
            json.dump({"code_verifier": code_verifier}, f)
    return auth_url


def exchange_auth_code(code: str, redirect_uri: str) -> None:
    """Exchange an OAuth2 authorization code for credentials and save the token."""
    from google_auth_oauthlib.flow import Flow

    flow = Flow.from_client_secrets_file(
        CREDENTIALS_FILE,
        scopes=SCOPES,
        redirect_uri=redirect_uri,
    )
    # Restore PKCE code_verifier if it was saved during get_auth_url
    if os.path.exists(CODE_VERIFIER_FILE):
        with open(CODE_VERIFIER_FILE, "r") as f:
            data = json.load(f)
        os.remove(CODE_VERIFIER_FILE)
        if data.get("code_verifier"):
            flow.code_verifier = data["code_verifier"]
    flow.fetch_token(code=code)
    creds = flow.credentials
    with open(TOKEN_FILE, "w") as f:
        f.write(creds.to_json())


def _build_service():
    from googleapiclient.discovery import build
    return build("drive", "v3", credentials=_get_credentials())


def is_authenticated() -> bool:
    """Check whether valid Drive credentials with the required scopes exist, without triggering auth flow."""
    try:
        from google.oauth2.credentials import Credentials
        from google.auth.transport.requests import Request

        if not os.path.exists(TOKEN_FILE) or not _token_has_required_scopes():
            return False
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
            with open(TOKEN_FILE, "w") as f:
                f.write(creds.to_json())
        return creds is not None and creds.valid
    except Exception:
        return False


def needs_reauth() -> bool:
    """True if a token exists but is missing newly-added scopes (e.g. drive.readonly)."""
    return os.path.exists(TOKEN_FILE) and not _token_has_required_scopes()


def has_credentials_file() -> bool:
    return os.path.exists(CREDENTIALS_FILE)


# ─────────────────────────────────────────────
# Upload & share
# ─────────────────────────────────────────────

_drive_id_cache: dict[str, str] = {}  # local_path → drive_file_id


def upload_and_get_url(local_path: str) -> str:
    """
    Upload a local file to Google Drive (if not already uploaded),
    make it publicly readable, and return a direct-download URL.
    """
    local_path = str(Path(local_path).resolve())

    # Return cached URL if already uploaded this session
    if local_path in _drive_id_cache:
        file_id = _drive_id_cache[local_path]
        return _direct_url(file_id)

    service = _build_service()
    filename = Path(local_path).name
    mime_type, _ = mimetypes.guess_type(local_path)
    mime_type = mime_type or "application/octet-stream"

    # Check if already uploaded (search by name + appProperty marker)
    results = service.files().list(
        q=f"name='{filename}' and trashed=false and appProperties has {{ key='localPath' and value='{filename}' }}",
        fields="files(id, name)",
    ).execute()

    existing = results.get("files", [])
    if existing:
        file_id = existing[0]["id"]
    else:
        # Upload
        from googleapiclient.http import MediaFileUpload
        file_metadata = {
            "name": filename,
            "appProperties": {"localPath": filename},
        }
        media = MediaFileUpload(local_path, mimetype=mime_type, resumable=True)
        uploaded = service.files().create(
            body=file_metadata,
            media_body=media,
            fields="id",
        ).execute()
        file_id = uploaded["id"]

    # Make publicly readable
    service.permissions().create(
        fileId=file_id,
        body={"type": "anyone", "role": "reader"},
    ).execute()

    _drive_id_cache[local_path] = file_id
    return _direct_url(file_id)


def _direct_url(file_id: str) -> str:
    return f"https://drive.google.com/uc?export=download&id={file_id}"


# ─────────────────────────────────────────────
# Browse arbitrary Drive folders by ID (Resource Library)
# ─────────────────────────────────────────────

import re

_FOLDER_ID_RE = re.compile(r"/folders/([a-zA-Z0-9_-]+)")
_GENERIC_ID_RE = re.compile(r"[?&]id=([a-zA-Z0-9_-]+)")


def extract_folder_id(url_or_id: str) -> str:
    """Pull a Drive folder/file ID out of a full Drive URL, or pass through if already an ID."""
    url_or_id = url_or_id.strip()
    m = _FOLDER_ID_RE.search(url_or_id)
    if m:
        return m.group(1)
    m = _GENERIC_ID_RE.search(url_or_id)
    if m:
        return m.group(1)
    # Assume it's already a bare ID
    return url_or_id


_FILE_FIELDS = "id, name, mimeType, webViewLink, iconLink, modifiedTime, size"


def get_folder_meta(folder_id: str) -> dict:
    """Get name/metadata for a single folder/file by ID."""
    service = _build_service()
    return service.files().get(
        fileId=folder_id,
        fields=_FILE_FIELDS,
        supportsAllDrives=True,
    ).execute()


def list_drive_folder(folder_id: str) -> dict:
    """
    List subfolders and files inside a Drive folder by its ID (via the Drive API,
    independent of the locally-synced folder).
    Returns { folders: [...], files: [...] }
    """
    service = _build_service()
    results = service.files().list(
        q=f"'{folder_id}' in parents and trashed=false",
        fields=f"files({_FILE_FIELDS})",
        orderBy="folder,name_natural",
        pageSize=300,
        supportsAllDrives=True,
        includeItemsFromAllDrives=True,
    ).execute()

    items = results.get("files", [])
    folders = [i for i in items if i["mimeType"] == GOOGLE_FOLDER_MIME]
    files = [i for i in items if i["mimeType"] != GOOGLE_FOLDER_MIME]

    for f in files:
        f["view_url"] = f.get("webViewLink") or f"https://drive.google.com/file/d/{f['id']}/view"
        if "size" in f:
            try:
                f["size"] = int(f["size"])
            except (TypeError, ValueError):
                pass
        mime = f.get("mimeType", "")
        f["media_type"] = (
            "image" if mime.startswith("image/") else
            "video" if mime.startswith("video/") else
            "document"
        )

    return {"folders": folders, "files": files}


def get_drive_direct_url(file_id: str) -> str:
    """
    Make an existing Drive file (already living in the Social Media folder)
    publicly readable and return a direct-download URL for it. No upload
    needed since the file is already on Drive.
    """
    service = _build_service()
    try:
        service.permissions().create(
            fileId=file_id,
            body={"type": "anyone", "role": "reader"},
            supportsAllDrives=True,
        ).execute()
    except Exception:
        pass  # already public
    return _direct_url(file_id)


def get_file_text(file_id: str, mime_type: str = None) -> str:
    """
    Extract plain text from a Drive file: Google Docs (export), PDFs (pypdf),
    and Word .docx (python-docx). Returns "" for unsupported types.
    """
    import io
    from googleapiclient.http import MediaIoBaseDownload

    service = _build_service()

    if mime_type is None:
        meta = service.files().get(fileId=file_id, fields="mimeType").execute()
        mime_type = meta.get("mimeType")

    if mime_type == GOOGLE_DOC_MIME:
        data = service.files().export(fileId=file_id, mimeType="text/plain").execute()
        return data.decode("utf-8") if isinstance(data, bytes) else str(data)

    if mime_type in (PDF_MIME, DOCX_MIME):
        request = service.files().get_media(fileId=file_id, supportsAllDrives=True)
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        buf.seek(0)

        if mime_type == PDF_MIME:
            from pypdf import PdfReader
            reader = PdfReader(buf)
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
            if len(text.strip()) < 30:
                # Likely a scanned/image-only PDF — OCR each page image isn't
                # attempted here (no poppler dependency); fall through empty.
                return text
            return text
        else:
            from docx import Document
            doc = Document(buf)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            text = "\n".join(p for p in parts if p)

            if len(text.strip()) < 30:
                # Likely a doc whose content is a pasted screenshot — OCR any
                # embedded images with Claude vision.
                text = _ocr_embedded_images(buf)

            return text


def _ocr_embedded_images(docx_buf) -> str:
    """Extract embedded images from a .docx (zip) and OCR them via Claude vision."""
    import zipfile
    from agent import transcribe_image_text

    docx_buf.seek(0)
    z = zipfile.ZipFile(docx_buf)
    media_files = sorted(
        n for n in z.namelist()
        if n.startswith("word/media/") and n.lower().endswith((".png", ".jpg", ".jpeg"))
    )

    parts = []
    for n in media_files:
        data = z.read(n)
        media_type = "image/jpeg" if n.lower().endswith((".jpg", ".jpeg")) else "image/png"
        try:
            parts.append(transcribe_image_text(data, media_type))
        except Exception as e:
            parts.append(f"[Could not read embedded image {n}: {e}]")

    return "\n\n".join(parts)

    return ""
